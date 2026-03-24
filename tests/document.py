from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Создаем новый документ
document = Document()

# ============= СОЗДАНИЕ ЗАГОЛОВКА =============

# Добавляем заголовок первого уровня
heading = document.add_heading('Отчет о продажах', level=1)

# Применяем стиль к заголовку
heading.style = document.styles['Title']  # Встроенный стиль Title

# Или можно настроить заголовок вручную
heading = document.add_heading('Анализ продаж', level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Выравнивание по центру

# Для каждого run в заголовке можно задать форматирование
for run in heading.runs:
    run.font.size = Pt(24)  # Размер шрифта
    run.font.bold = True  # Жирный
    run.font.color.rgb = RGBColor(0, 51, 102)  # Цвет

# Добавляем подзаголовок
subheading = document.add_heading('Данные за декабрь 2024', level=2)
subheading.style = document.styles['Heading 2']

# ============= СОЗДАНИЕ ТАБЛИЦЫ =============

# Способ 1: Создание таблицы с указанием количества строк и столбцов
table = document.add_table(rows=4, cols=4)
table.style = 'Table Grid'  # Применяем стиль таблицы
table.alignment = WD_TABLE_ALIGNMENT.CENTER  # Выравнивание таблицы по центру

# Заполняем таблицу данными
# Заголовки таблицы
headers = ['Продукт', 'Количество', 'Цена', 'Сумма']
data = [
    ['Ноутбук', '5', '50000', '250000'],
    ['Мышь', '10', '1500', '15000'],
    ['Клавиатура', '7', '3000', '21000']
]

# Заполняем заголовки
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    # Применяем стиль к ячейке заголовка
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(12)
    # Заливка фона для заголовков
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'D9E1F2')
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Заполняем данные
for i, row_data in enumerate(data, start=1):
    for j, value in enumerate(row_data):
        cell = table.rows[i].cells[j]
        cell.text = value
        # Выравнивание текста в ячейках
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============= СОЗДАНИЕ ТАБЛИЦЫ С КАСТОМНОЙ ШИРИНОЙ =============

# Создаем таблицу с кастомной шириной столбцов
table2 = document.add_table(rows=3, cols=3)
table2.style = 'Light Shading Accent 1'

# Устанавливаем ширину столбцов
widths = [Inches(2), Inches(1.5), Inches(2.5)]
for row in table2.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = width

# Добавляем данные
table2.rows[0].cells[0].text = 'Категория'
table2.rows[0].cells[1].text = 'Количество'
table2.rows[0].cells[2].text = 'Процент'

table2.rows[1].cells[0].text = 'Электроника'
table2.rows[1].cells[1].text = '45'
table2.rows[1].cells[2].text = '60%'

table2.rows[2].cells[0].text = 'Аксессуары'
table2.rows[2].cells[1].text = '30'
table2.rows[2].cells[2].text = '40%'

# ============= ПРИМЕНЕНИЕ СТИЛЕЙ К ТАБЛИЦЕ =============

# Создаем таблицу с применением различных стилей
# Доступные стили: 'Table Grid', 'Light Shading', 'Light Shading Accent 1-6',
# 'Light List', 'Light List Accent 1-6', 'Medium Shading 1-2' и др.

# Применяем стиль к существующей таблице
table3 = document.add_table(rows=2, cols=2)
table3.style = 'Medium Shading 2 Accent 1'

table3.rows[0].cells[0].text = 'Показатель'
table3.rows[0].cells[1].text = 'Значение'
table3.rows[1].cells[0].text = 'Выручка'
table3.rows[1].cells[1].text = '286 000 руб.'

# ============= РАСШИРЕННОЕ ФОРМАТИРОВАНИЕ ТАБЛИЦЫ =============

def format_table_cell(cell, text, bold=False, font_size=11,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT, bg_color=None):
    """Функция для форматирования ячейки таблицы"""
    cell.text = text
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment

    run = paragraph.runs[0]
    run.font.size = Pt(font_size)
    run.font.bold = bold

    if bg_color:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), bg_color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

# Создаем таблицу с форматированием
table4 = document.add_table(rows=4, cols=3)
table4.style = 'Table Grid'

# Форматируем заголовки
headers_with_style = ['Месяц', 'План', 'Факт']
for i, header in enumerate(headers_with_style):
    format_table_cell(table4.rows[0].cells[i], header,
                     bold=True, font_size=12,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER,
                     bg_color='E7E6E6')

# Форматируем данные
data_with_style = [
    ['Январь', '100000', '95000'],
    ['Февраль', '110000', '112000'],
    ['Март', '120000', '125000']
]

for i, row_data in enumerate(data_with_style, start=1):
    for j, value in enumerate(row_data):
        alignment = WD_ALIGN_PARAGRAPH.RIGHT if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
        format_table_cell(table4.rows[i].cells[j], value,
                         font_size=11, alignment=alignment)

# ============= ОБЪЕДИНЕНИЕ ЯЧЕЕК =============

table5 = document.add_table(rows=3, cols=3)
table5.style = 'Table Grid'

# Объединяем ячейки в первой строке
cell = table5.cell(0, 0)
cell.merge(table5.cell(0, 2))
cell.text = 'Объединенная ячейка'
cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Заполняем остальные ячейки
table5.rows[1].cells[0].text = 'A1'
table5.rows[1].cells[1].text = 'B1'
table5.rows[1].cells[2].text = 'C1'
table5.rows[2].cells[0].text = 'A2'
table5.rows[2].cells[1].text = 'B2'
table5.rows[2].cells[2].text = 'C2'

# ============= ДОБАВЛЕНИЕ ТАБЛИЦЫ ИЗ СПИСКА СЛОВАРЕЙ =============

def create_table_from_dict(document, data, headers):
    """Создание таблицы из списка словарей"""
    if not data:
        return None

    # Создаем таблицу с заголовками
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 1'

    # Заполняем заголовки
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    # Добавляем данные
    for item in data:
        row_cells = table.add_row().cells
        for i, header in enumerate(headers):
            row_cells[i].text = str(item.get(header, ''))

    return table

# Пример использования
sales_data = [
    {'Продукт': 'Ноутбук', 'Продажи': 25, 'Выручка': 1250000},
    {'Продукт': 'Телефон', 'Продажи': 40, 'Выручка': 800000},
    {'Продукт': 'Планшет', 'Продажи': 15, 'Выручка': 450000}
]

headers = ['Продукт', 'Продажи', 'Выручка']
create_table_from_dict(document, sales_data, headers)

# Сохраняем документ
document.save('report_with_tables.docx')
print("Документ успешно создан!")