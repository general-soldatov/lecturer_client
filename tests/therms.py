from docx import Document
from docx.shared import Inches, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_table_with_one_autofit_cell(font_size = 1):
    doc = Document()

    # Создаем таблицу с 3 строками и 3 колонками
    # table = doc.add_table(rows=4, cols=3)
    # table.style = 'Table Grid'
    # table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # # Отключаем автоподбор для всей таблицы
    # table.autofit = False
    # table.width = Cm(16)  # Ширина всей таблицы на всю страницу

    # Устанавливаем ширину для всех колонок, кроме той, что хотим "по содержимому"
    # Допустим, 2-я колонка (индекс 1) должна быть по содержимому
    # table.columns[0].width = Inches(2.0)   # Фиксированная ширина
    # table.columns[2].width = Inches(2.0)   # Фиксированная ширина

    # Для колонки, которую хотим по содержимому, НЕ устанавливаем ширину

    # Заполняем заголовки
    # headers = ["Столбец 1 (фикс)", "Столбец 2 (авто)", "Столбец 3 (фикс)"]
    # for i, header in enumerate(headers):
    #     table.rows[0].cells[i].text = header

    # Данные разной длины для демонстрации
    data = [
        ["Коротко", "Очень-очень длинный текст для демонстрации авто-подбора ширины", "OK"],
        ["A", "B", "C"],
        ["Кратко", "Экстремально длинное предложение, которое должно заставить колонку расшириться автоматически", "D"]
    ]

    for i, row_data in enumerate(data, start=1):
        table = doc.add_table(rows=4, cols=2)
        table.columns[1].width = Inches(6.1 - font_size / 72 * len(row_data[0]))
        # table.style = 'Table Grid'
        # table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Отключаем автоподбор для всей таблицы
        # table.autofit = False
        table.width = Cm(16)
        # table.columns[0].width = Inches(font_size / 72 * len(row_data[0]))
        table.rows[0].cells[0].text = row_data[0]

        # КРИТИЧЕСКИ ВАЖНО: Принудительно сбрасываем ширину для ячеек авто-колонки
        for row in table.rows:
            # Для ячейки в авто-колонке (индекс 1) удаляем явно заданную ширину
            # for cell in row.cells:
            cell = row.cells[0]
            tcPr = cell._tc.get_or_add_tcPr()

            # Удаляем элемент ширины, если он есть
            for w in tcPr.findall(qn('w:tcW')):
                tcPr.remove(w)

            # Устанавливаем тип ширины как "auto"
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:type'), 'auto')
            tcPr.append(tcW)

        doc.add_paragraph()

    doc.save('table_mixed_width.docx')
    print("✅ Таблица с одной авто-колонкой создана!")

# create_table_with_one_autofit_cell()

from docx import Document
from docx.shared import Inches, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_adaptive_table():
    """Создает таблицу: 1-й столбец по содержимому, 2-й - на всю страницу"""
    doc = Document()

    # Параметры страницы A4 (ширина ~16.5 см с учетом полей)
    page_width = Cm(16)  # Доступная ширина страницы

    # Создаем таблицу с 2 столбцами

    # Заполняем таблицу данными разной длины
    data = [
        ("Коротко", "Это очень длинный текст, который будет занимать всю ширину страницы и переноситься на следующие строки автоматически"),
        ("Средний текст", "Вторая колонка растягивается на всю доступную ширину, независимо от содержимого первой колонки"),
        ("Очень-очень-очень длинное слово в первой колонке", "Первая колонка подстраивается под длину слова, а вторая занимает всё остальное пространство"),
        ("Кратко", "Это пример таблицы с двумя колонками разного типа"),
        ("Тест", "Первая колонка может быть любой ширины в зависимости от самого длинного слова в ней")
    ]

    # Заголовки
    # table.rows[0].cells[0].text = "Термин"
    # table.rows[0].cells[1].text = "Определение"

    # # Делаем заголовки жирными
    # for i in range(2):
    #     table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    # Заполняем данные
    for i, (term, definition) in enumerate(data, start=1):
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Отключаем автоподбор для ручного управления
        table.autofit = True

        # Устанавливаем общую ширину таблицы на всю страницу
        table.width = page_width

        # 1. Настраиваем первый столбец на авто-ширину (по содержимому)
        # Для этого НЕ устанавливаем ему фиксированную ширину

        # 2. Второй столбец заполняет оставшееся пространство
        # Устанавливаем ему большую ширину
        table.columns[0].width = Cm(5)
        table.columns[1].width = page_width
        # КРИТИЧЕСКИ ВАЖНО: Очищаем ширину для ячеек первого столбца
        # for row in table.rows:
        #     # Для первого столбца удаляем явную ширину
        #     cell = row.cells[0]
        #     tcPr = cell._tc.get_or_add_tcPr()

        #     # Удаляем существующие настройки ширины
        #     for w in tcPr.findall(qn('w:tcW')):
        #         tcPr.remove(w)

        #     # Устанавливаем тип "auto" для автоматического подбора
        #     tcW = OxmlElement('w:tcW')
        #     tcW.set(qn('w:type'), 'auto')
        #     tcPr.append(tcW)

            # Для второго столбца устанавливаем оставшуюся ширину
        # cell2 = row.cells[1]
        # cell2.width = page_width

        table.rows[0].cells[0].text = term
        table.rows[0].cells[1].text = ''
        doc.add_paragraph()

    doc.save("adaptive_two_columns.docx")
    print("✅ Таблица создана: 1-й столбец по содержимому, 2-й на всю страницу")

# create_adaptive_table()
import docx
from docx.shared import Inches, Pt
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.enum.table import WD_ROW_HEIGHT_RULE

class ThermsTask:
    def __init__(self, doc: docx.Document, rows = 4, font_size=12, font_name="Times New Roman"):
        self.doc = doc
        self.font_size = font_size
        self.font_name = font_name
        self.rows = rows

    @staticmethod
    def set_cell_bottom_border(cell, color="000000", sz="4", val="single"):
        """
        Добавляет нижнюю границу к ячейке.
        color: HEX-код цвета (например, "000000" - черный, "CCCCCC" - серый)
        sz: размер/толщина границы (4 = 0.5 pt, 8 = 1 pt, 12 = 1.5 pt)
        val: тип линии (single, double, dashed и т.д.)
        """
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # Создаем элемент нижней границы
        bottom = parse_xml(f'<w:bottom {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>')
        tcBorders.append(bottom)

    @staticmethod
    def set_cell_text_with_font(cell, text, font_name="Times New Roman", font_size=11):
        """Очищает ячейку, добавляет текст и применяет к нему шрифт."""
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        p.text = ""  # Очищаем дефолтный текст
        # Добавляем текст через run для форматирования
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)

    def add_table(self, word: str, coeff=0.6):
        table = self.doc.add_table(rows=self.rows, cols=2)

        # Отключаем автоматический подбор, чтобы задать ширину вручную
        table.allow_autofit = False

        PAGE_WIDTH = Inches(6.25)
        col_1_width = Inches(len(word) * 1/72 * self.font_size * coeff)
        for cell in table.columns[0].cells:
            cell.width = col_1_width
        col_2_width = PAGE_WIDTH - col_1_width
        for cell in table.columns[1].cells:
            cell.width = col_2_width
        for i, row in enumerate(table.rows):
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Pt(1.3 * self.font_size)
            if i == 0:
                self.set_cell_text_with_font(row.cells[0], word,
                                             font_size=self.font_size,
                                             font_name=self.font_name)
                self.set_cell_bottom_border(row.cells[1], color="000000", sz="6")
                continue
            row.cells[0].merge(row.cells[1])
            self.set_cell_bottom_border(row.cells[0], color="000000", sz="6")

        self.doc.add_paragraph()

    def create_tables(self, therms: list):
        for therm in therms:
            self.add_table(therm)

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls

def set_cell_border(cell, **kwargs):
    """
    Универсальная функция для настройки границ ячейки.
    Использование: set_cell_border(cell, right={"sz": 6, "color": "000000", "val": "single"})
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            b_el = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="{edge_data.get("val", "single")}" '
                             f'w:sz="{edge_data.get("sz", 4)}" w:space="0" '
                             f'w:color="{edge_data.get("color", "auto")}"/>')
            tcBorders.append(b_el)

def add_styled_paragraph(doc_or_cell, text, font_name="Arial", font_size=11, bold=False, space_after=6):
    """Добавляет абзац с заданным шрифтом и отступами."""
    p = doc_or_cell.add_paragraph() if hasattr(doc_or_cell, 'add_paragraph') else doc_or_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)

    if text:
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
    return p

def create_physics_task(task_text, image_path=None, data_given=None, data_find=None):
    doc = Document()

    # Расчет рабочей ширины страницы
    section = doc.sections[0]
    available_width = section.page_width - section.left_margin - section.right_margin

    # 1. Текст задачи
    p_task = add_styled_paragraph(doc, f"Задача. {task_text}", font_name="Arial", font_size=12, bold=False)
    p_task.runs[0].font.bold = True  # Делаем слово "Задача." жирным
    p_task.runs[0].text = "Задача. "
    p_task.add_run(task_text)

    # 2. Вставка рисунка (если файл существует)
    if image_path and os.path.exists(image_path):
        p_img = doc.add_paragraph()
        p_img.paragraph_format.space_after = Pt(12)
        p_img.paragraph_format.space_before = Pt(6)
        run_img = p_img.add_run()
        # Ограничиваем рисунок максимум 4 дюймами в ширину, чтобы не сломать верстку
        run_img.add_picture(image_path, width=Inches(4.0))

    # 3. Создание таблицы «Дано / Решение»
    # Строка 1: Дано | Решение
    # Строка 2: Найти | Решение (продолжение)
    table = doc.add_table(rows=2, cols=2)
    table.autofit = True

    # Включаем XML автоподбор
    tblPr = table._tbl.tblPr
    tblLayout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="autofit"/>')
    tblPr.append(tblLayout)

    # Настраиваем размеры: первый столбец сжимается под контент, второй — на всю страницу
    for row in table.rows:
        row.cells[0].width = Inches(0)
        row.cells[1].width = available_width
        row.height = Pt(40)  # Минимальная базовая высота для аккуратного вида
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    # Заполнение блока "ДАНО" (Строка 0, Колонка 0)
    cell_given = table.cell(0, 0)
    add_styled_paragraph(cell_given, "Дано:", bold=True)
    for key, value in data_given.items():
        add_styled_paragraph(cell_given, f"{key} = {value}")

    # Заполнение блока "НАЙТИ" (Строка 1, Колонка 0)
    cell_find = table.cell(1, 0)
    add_styled_paragraph(cell_find, "Найти:", bold=True)
    for item in data_find:
        add_styled_paragraph(cell_find, f"{item} — ?")

    # Заполнение блока "РЕШЕНИЕ" (Строка 0, Колонка 1)
    cell_sol = table.cell(0, 1)
    add_styled_paragraph(cell_sol, "Решение:", bold=True)
    add_styled_paragraph(cell_sol, "1. Согласно второму закону Ньютона: F = m * a")
    add_styled_paragraph(cell_sol, "2. Выразим ускорение: a = F / m")
    add_styled_paragraph(cell_sol, "3. Подставим значения и найдем ответ.")

    # 4. Отрисовка разделительных линий физического шаблона
    # Вертикальная линия: справа от "Дано" и "Найти"
    border_format = {"sz": 6, "color": "404040", "val": "single"} # 0.75 pt, темно-серый
    set_cell_border(table.cell(0, 0), right=border_format)
    set_cell_border(table.cell(1, 0), right=border_format)

    # Горизонтальная линия: снизу ячейки "Дано" (между Дано и Найти)
    set_cell_border(table.cell(0, 0), bottom=border_format, right=border_format)

    # Сохранение документа
    doc.save("physics_task_template.docx")

import math2docx
from docx import Document

def add_to_docx_with_latex(document, text):
    """
    Создает новый .docx файл и вставляет в него формулы из LaTeX.
    """
    paragraph = document.add_paragraph()
    for i, elem in enumerate(text.split('$')):
        if i % 2:
            # Ключевая функция: math2docx.add_math добавляет формулу в абзац
            math2docx.add_math(paragraph, elem)
        else:
            paragraph.add_run(elem)




# --- Пример использования ---
if __name__ == "__main__":
    document = Document()
    document.add_heading('Документ с формулами из LaTeX', level=1)

    # 2. Определяем формулы, которые хотим вставить
    formulas = [
        (r"Квадратное уравнение: $x = \frac{-b \pm \sqrt{b^2 - 4ac}}{2a}$ example of formula to find $x$"),
        (r"Формула Эйлера: $e^{i\pi} + 1 = 0$"),
        (r"Определенный интеграл: $\int_{0}^{\infty} e^{-x^2} dx = \frac{\sqrt{\pi}}{2}$")
    ]

    # 3. Вставляем каждую формулу в новый абзац
    output_path = "formulas_example.docx"
    for formula in formulas:
        add_to_docx_with_latex(document, formula)
    # 4. Сохраняем документ
    document.save(output_path)
    print(f"✅ Документ '{output_path}' успешно создан с формулами!")