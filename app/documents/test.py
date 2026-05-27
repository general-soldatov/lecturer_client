import os
import requests
import math2docx
from random import shuffle
from docx import Document
from docx.document import Document as Docx
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from enum import Enum
from typing import List, Tuple, Optional

from app.documents.test_make import Project, QuestionsType


class QuestionType(Enum):
    """Типы вопросов"""
    SINGLE_CHOICE = "single_choice"      # Один правильный ответ
    MULTIPLE_CHOICE = "multiple_choice"  # Несколько правильных ответов
    MATCHING = "matching"                # Сопоставление
    SORTING = "sorting"                  # Сортировка
    NUMBER = "number"                    # Числовая задача

class ThermsTask:
    def __init__(self, doc: Docx, rows = 4, font_size=12, font_name="Times New Roman"):
        self.doc = doc
        self.font_size = font_size
        self.font_name = font_name
        self.rows = rows

    @staticmethod
    def set_cell_bottom_border(cell, color="000000", sz="4", val="single"):
        """
        Добавляет нижнюю границу к ячейке.
        color: HEX-код цвета (например, "000000" - черный, "CCCCCC" - серый)
        sz: размер/толщина границы (4 = 0.5 pt, 8 = 1 pt, 12 = 1.5 pt)
        val: тип линии (single, double, dashed и т.д.)
        """
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # Создаем элемент нижней границы
        bottom = parse_xml(f'<w:bottom {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>')
        tcBorders.append(bottom)

    @staticmethod
    def set_cell_text_with_font(cell, text, font_name="Times New Roman", font_size=11):
        """Очищает ячейку, добавляет текст и применяет к нему шрифт."""
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        p.text = ""  # Очищаем дефолтный текст
        # Добавляем текст через run для форматирования
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)

    def add_table(self, word: str, coeff=0.6):
        table = self.doc.add_table(rows=self.rows, cols=2)

        # Отключаем автоматический подбор, чтобы задать ширину вручную
        table.allow_autofit = False

        PAGE_WIDTH = Inches(6.25)
        col_1_width = Inches(len(word) * 1/72 * self.font_size * coeff)
        for cell in table.columns[0].cells:
            cell.width = col_1_width
        col_2_width = PAGE_WIDTH - col_1_width
        for cell in table.columns[1].cells:
            cell.width = col_2_width
        for i, row in enumerate(table.rows):
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Pt(1.3 * self.font_size)
            if i == 0:
                self.set_cell_text_with_font(row.cells[0], word,
                                             font_size=self.font_size,
                                             font_name=self.font_name)
                self.set_cell_bottom_border(row.cells[1], color="000000", sz="6")
                continue
            row.cells[0].merge(row.cells[1])
            self.set_cell_bottom_border(row.cells[0], color="000000", sz="6")

        par = self.doc.add_paragraph()
        par.paragraph_format.line_spacing = 0

    def create_tables(self, therms: list):
        for therm in therms:
            self.add_table(therm)

class TestQuestion:
    """Базовый класс для тестового вопроса"""

    def __init__(self, question_text: str, points: int = 1):
        self.question_text = question_text
        self.points = points
        self.question_type = None

    @staticmethod
    def set_color_cell(cell, color='#FFFFFF'):
        tcPr = cell._tc.get_or_add_tcPr()

        # Удаляем существующее форматирование
        for shading in tcPr.findall(qn('w:shd')):
            tcPr.remove(shading)

        # Создаем новую заливку с белым цветом
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:fill'), color[1:])
        tcPr.append(shading)

    @staticmethod
    def add_to_docx_with_latex(paragraph, text, bold: bool = False):
        """
        Создает новый .docx файл и вставляет в него формулы из LaTeX.
        """
        for i, elem in enumerate(text.split('$')):
            if i % 2:
                # Ключевая функция: math2docx.add_math добавляет формулу в абзац
                math2docx.add_math(paragraph, elem)
            else:
                paragraph.add_run(elem).bold = bold

    def _apply_cell_formatting(self, cell, text: str, font_size: int = 11,
                               bold: bool = False, italic: bool = False,
                               alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
                               bg_color: str = None):
        """Применяет форматирование к ячейке таблицы"""
        cell.text = text
        paragraph = cell.paragraphs[0]
        paragraph.alignment = alignment

        if paragraph.runs:
            run = paragraph.runs[0]
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.italic = italic

        if bg_color:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), bg_color)
            cell._tc.get_or_add_tcPr().append(shading_elm)

    def _add_numbering(self, index: int) -> str:
        """Возвращает букву для нумерации вариантов"""
        return f"{chr(65 + index)}."

    def _set_cell_border(self, cell, border_color: str = "000000", border_size: int = 4):
        """Устанавливает границы ячейки"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(border_size))
            border.set(qn('w:color'), border_color)
            tcPr.append(border)

    def set_cell_border(self, cell, **kwargs):
        """
        Универсальная функция для настройки границ ячейки.
        Использование: set_cell_border(cell, right={"sz": 6, "color": "000000", "val": "single"})
        """
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                b_el = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="{edge_data.get("val", "single")}" '
                                f'w:sz="{edge_data.get("sz", 4)}" w:space="0" '
                                f'w:color="{edge_data.get("color", "auto")}"/>')
                tcBorders.append(b_el)

    def add_styled_paragraph(self, doc: Optional[Docx] = None, cell = None, text = None, font_name="Arial", font_size=11, bold=False, space_after=6):
        """Добавляет абзац с заданным шрифтом и отступами."""
        p = doc
        if cell:
            p = cell.paragraphs[0]
        # p = doc_or_cell.add_paragraph() if hasattr(doc_or_cell, 'add_paragraph') else doc_or_cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)

        if text:
            run = p.add_run(text)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
        return p

    def add_header_question(self, doc: Document, question_number: int, description = None):
        # Создаем основную таблицу для вопроса
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Light Shading'
        table.width = Cm(15.25)

        # Настраиваем ячейку вопроса
        question_cell = table.cell(0, 0)
        question_cell.width = Cm(15.25)

        # Добавляем текст вопроса
        question_para = question_cell.paragraphs[0]
        question_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        question_para.add_run(f"Вопрос {question_number}. ").bold = True
        if description:
            question_para.add_run(f"({description}) ").bold = True
        question_para.add_run(self.question_text).bold = False
        # question_para.add_run(f" [{self.points} балл]").italic = True
        question_para.paragraph_format.space_after = Pt(6)

        return question_cell


class SingleChoiceQuestion(TestQuestion):
    """Класс для вопроса с одним правильным ответом"""

    def __init__(self, question_text: str, options: List[str], correct_answer: int, points: int = 1):
        super().__init__(question_text, points)
        self.options = options
        self.correct_answer = correct_answer
        self.question_type = QuestionType.SINGLE_CHOICE

    def render_to_document(self, doc: Document, question_number: int):
        """Отрисовывает вопрос в документе"""
        question_cell = self.add_header_question(doc, question_number, "")
        # Добавляем варианты ответов
        for i, option in enumerate(self.options):
            option_para = question_cell.add_paragraph()
            option_para.bold = True
            option_para.add_run("☐ ").font.size = Pt(12)
            option_para.add_run(f"{self._add_numbering(i)} ").bold = False
            option_para.add_run(option).bold = False
            option_para.paragraph_format.left_indent = Pt(20)

            # Для демонстрации помечаем правильный ответ (в реальном тесте это скрыто)
            # if i == self.correct_answer:
            #     run = option_para.add_run(" ✓")
            #     run.font.color.rgb = RGBColor(0, 128, 0)
            #     run.font.bold = True

        # Добавляем поле для ответа
        # answer_para = question_cell.add_paragraph()
        # answer_para.add_run("Ответ: __________").italic = True
        # answer_para.paragraph_format.space_before = Pt(6)

        # Добавляем отступ после вопроса
        # doc.add_paragraph()


class MultipleChoiceQuestion(TestQuestion):
    """Класс для вопроса с несколькими правильными ответами"""

    def __init__(self, question_text: str, options: List[str], correct_answers: List[int], points: int = 2):
        super().__init__(question_text, points)
        self.options = options
        self.correct_answers = correct_answers
        self.question_type = QuestionType.MULTIPLE_CHOICE

    def render_to_document(self, doc: Document, question_number: int):
        """Отрисовывает вопрос в документе"""
        # table = doc.add_table(rows=1, cols=1)
        # table.style = 'Light Shading'
        # # table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # table.width = Cm(16)

        # question_cell = table.cell(0, 0)

        # # Заголовок вопроса
        # question_para = question_cell.paragraphs[0]
        # question_para.add_run(f"Вопрос {question_number} (Выберите несколько вариантов). ").bold = True
        # question_para.add_run(self.question_text).bold = False
        # # question_para.add_run(f" [{self.points} балл]").italic = True
        # question_para.paragraph_format.space_after = Pt(6)
        question_cell = self.add_header_question(doc, question_number, "Выберите несколько вариантов")
        # Варианты ответов с чекбоксами
        for i, option in enumerate(self.options):
            option_para = question_cell.add_paragraph()
            option_para.bold = True
            option_para.add_run("☐ ").font.size = Pt(12)
            option_para.add_run(f"{self._add_numbering(i)} ").bold = False
            option_para.add_run(option).bold = False
            option_para.paragraph_format.left_indent = Pt(20)


            # Помечаем правильные ответы (для демонстрации)
            # if i in self.correct_answers:
            #     run = option_para.add_run(" ✓")
            #     run.font.color.rgb = RGBColor(0, 128, 0)
            #     run.font.bold = True

        # doc.add_paragraph()


class MatchingQuestion(TestQuestion):
    """Класс для вопроса на сопоставление"""

    def __init__(self, question_text: str, left_items: List[str], right_items: List[str],
                 matches: List[Tuple[int, int]], points: int = 3):
        super().__init__(question_text, points)
        self.left_items = left_items
        self.right_items = right_items
        self.matches = matches  # Список пар (индекс левого, индекс правого)
        self.question_type = QuestionType.MATCHING

    def render_to_document(self, doc: Document, question_number: int):
        """Отрисовывает вопрос в документе"""
        # Заголовок вопроса
        title_table = doc.add_table(rows=1, cols=1)
        title_table.style = 'Light Shading'
        title_table.width = Cm(16)

        title_cell = title_table.cell(0, 0)
        title_para = title_cell.paragraphs[0]
        title_para.add_run(f"Вопрос {question_number}. ").bold = True
        title_para.add_run(self.question_text).bold = False
        # title_para.add_run(f" [{self.points} балл]").italic = True

        # Создаем таблицу для сопоставления
        matching_table = doc.add_table(rows=max(len(self.left_items), len(self.right_items)), cols=2)
        # matching_table.style = 'Light Shading'
        # matching_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        matching_table.width = Cm(16)

        # Заголовки таблицы
        # headers = ['№', 'Элемент 1', 'Элемент 2']
        # for i, header in enumerate(headers):
        #     cell = matching_table.rows[0].cells[i]
        #     self._apply_cell_formatting(cell, header, font_size=12, bold=True,
        #                                alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Заполняем левую колонку
        for i, item in enumerate(self.left_items, start=0):
            # cell = matching_table.rows[i].cells[0]
            # self._apply_cell_formatting(cell, str(i), alignment=WD_ALIGN_PARAGRAPH.CENTER)

            cell = matching_table.rows[i].cells[0]
            self._apply_cell_formatting(cell, f"{self._add_numbering(i)} {item}")

        # Заполняем правую колонку
        for i, item in enumerate(self.right_items, start=0):
            cell = matching_table.rows[i].cells[1]
            self._apply_cell_formatting(cell, f"{i+1}. {item}")

        # Добавляем строки для ответов
        answer_table = doc.add_table(rows=3, cols=len(self.left_items))
        answer_table.style = 'Table Grid'
        answer_table.width = Cm(16)

        answer_cell = answer_table.rows[0].cells[0]
        self._apply_cell_formatting(answer_cell, "Ваши ответы:", font_size=11, bold=True)
        for i in range(1, max(len(self.left_items), len(self.right_items))):
            answer_cell.merge(answer_table.rows[0].cells[i])
        for i in range(max(len(self.left_items), len(self.right_items))):
            up_row = answer_table.rows[1].cells[i]
            self._apply_cell_formatting(up_row, self._add_numbering(i), alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=13, bold=True)
        answer_table.rows[2].height = Cm(0.75)

        # for i in range(len(self.left_items)):
        #     left_cell = answer_table.rows[i + 1].cells[0]
        #     self._apply_cell_formatting(left_cell, f"{i + 1}. → _____", alignment=WD_ALIGN_PARAGRAPH.LEFT)

        #     right_cell = answer_table.rows[i + 1].cells[1]
        #     self._apply_cell_formatting(right_cell, "", alignment=WD_ALIGN_PARAGRAPH.LEFT)

        # Для демонстрации показываем правильные ответы
        # demo_para = doc.add_paragraph()
        # demo_para.add_run("Пример правильных ответов: ").italic = True
        # for left_idx, right_idx in self.matches:
        #     demo_para.add_run(f"{left_idx + 1}→{right_idx + 1} ").font.size = Pt(9)

        # doc.add_paragraph()


class SortingQuestion(TestQuestion):
    """Класс для вопроса на сортировку"""

    def __init__(self, question_text: str, items: List[str], correct_order: List[int], points: int = 2):
        super().__init__(question_text, points)
        self.items = items
        self.correct_order = correct_order
        self.question_type = QuestionType.SORTING

    def render_to_document(self, doc: Document, question_number: int):
        """Отрисовывает вопрос в документе"""
        table = doc.add_table(rows=2, cols=1)
        table.style = 'Light Shading'
        # table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.width = Cm(16)

        # Заголовок вопроса
        header_cell = table.cell(0, 0)
        header_para = header_cell.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        header_para.add_run(f"Вопрос {question_number} (Расставьте в правильном порядке). ").bold = True
        header_para.add_run(self.question_text).bold = False
        # header_para.add_run(f" [{self.points} балл]").italic = True

        # Контент для сортировки
        content_cell = table.cell(1, 0)
        self.set_color_cell(content_cell)
        # content_cell.shading.background_color = RGBColor(255, 255, 255)
        # tcPr = content_cell._tc.get_or_add_tcPr()

        # # Удаляем существующее форматирование
        # for shading in tcPr.findall(qn('w:shd')):
        #     tcPr.remove(shading)

        # # Создаем новую заливку с белым цветом
        # shading = OxmlElement('w:shd')
        # shading.set(qn('w:val'), 'clear')
        # shading.set(qn('w:fill'), 'FFFFFF')
        # tcPr.append(shading)

        # Создаем пронумерованный список для сортировки
        for i, item in enumerate(self.items, start=1):
            item_para = content_cell.add_paragraph()
            item_para.add_run(f"{i}. ").bold = True
            item_para.add_run(item)
            item_para.paragraph_format.left_indent = Pt(20)
            item_para.paragraph_format.background_color = RGBColor(255, 255, 255)
            item_para.paragraph_format.space_after = Pt(3)

        # Добавляем поле для порядка
        order_para = content_cell.add_paragraph()
        order_para.add_run("\nПравильный порядок: ").italic = True
        order_para.add_run("___ → " * (len(self.items) - 1) + "___")
        order_para.paragraph_format.space_before = Pt(12)

        # Для демонстрации показываем правильный порядок
        # correct_order_str = " → ".join([str(i + 1) for i in self.correct_order])
        # demo_para = doc.add_paragraph()
        # demo_para.add_run(f"Пример: {correct_order_str}").italic = True
        # demo_para.runs[0].font.size = Pt(9)
        # demo_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        # doc.add_paragraph()
class NumberTask(TestQuestion):
    def __init__(self, question_text: str, image_path: Optional[str] = None, points: int = 3):
        super().__init__(question_text, points)
        self.image_path = image_path
        self.question_type = QuestionType.NUMBER

    def render_to_document(self, doc: Document, question_number: int):
        """Отрисовывает вопрос в документе"""
        table = doc.add_table(rows=2, cols=1)
        table.style = 'Light Shading'
        table.width = Cm(16)

        # Заголовок вопроса
        header_cell = table.cell(0, 0)
        header_para = header_cell.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        header_para.add_run(f"Вопрос {question_number} (Найдите численное решение). ").bold = True
        self.add_to_docx_with_latex(header_para, self.question_text)
        # Вставка рисунка (если файл существует)
        if self.image_path and os.path.exists(self.image_path):
            p_img = header_cell.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_after = Pt(12)
            p_img.paragraph_format.space_before = Pt(6)
            run_img = p_img.add_run()
            # Ограничиваем рисунок максимум 1.5 дюймами в высоту, чтобы не сломать верстку
            run_img.add_picture(self.image_path, height=Inches(1.5))
        last_cell = table.cell(1, 0)
        self.set_color_cell(last_cell)
        table_data = last_cell.add_table(rows=2, cols=2)
        table_data.autofit = True

        section = doc.sections[0]
        available_width = section.page_width - section.left_margin - section.right_margin

        # Включаем XML автоподбор
        tblPr = table_data._tbl.tblPr
        tblLayout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="autofit"/>')
        tblPr.append(tblLayout)

        # Настраиваем размеры: первый столбец сжимается под контент, второй — на всю страницу
        for row, size in zip(table_data.rows, [70, 40]):
            row.cells[0].width = Inches(2)
            row.cells[1].width = available_width
            row.height = Pt(size)  # Минимальная базовая высота для аккуратного вида
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        # Заполнение блока "ДАНО" (Строка 0, Колонка 0)
        for count, test in enumerate(["Дано:", "Решение:", "Найти:"]):
            cell_given = table_data.cell((count >> 1 & 1), (count >> 0 & 1))
            self.add_styled_paragraph(cell=cell_given, text=test, bold=True, font_name="Times New Roman")

        border_format = {"sz": 6, "color": "404040", "val": "single"} # 0.75 pt, темно-серый
        self.set_cell_border(table_data.cell(0, 0), right=border_format)
        self.set_cell_border(table_data.cell(1, 0), right=border_format)
        # Горизонтальная линия: снизу ячейки "Дано" (между Дано и Найти)
        self.set_cell_border(table_data.cell(0, 0), bottom=border_format, right=border_format)
        last_cell.add_paragraph()


class TestGenerator:
    """Генератор тестовых заданий"""

    def __init__(self, title: str, description: str = ""):
        self.document = Document()
        self.questions = []
        self.title = title
        self.description = description

    def add_question(self, question: TestQuestion):
        """Добавляет вопрос в тест"""
        self.questions.append(question)

    def add_header(self):
        """Добавляет заголовок теста"""
        # Основной заголовок
        title = self.document.add_heading(self.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(0, 51, 102)

        # Информационная таблица
        info_table = self.document.add_table(rows=2, cols=2)
        info_table.style = 'Table Grid'
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_table.width = Cm(16)

        info_data = [
            ("📋 Тип теста:", "Комплексное тестирование"),
            ("📝 Всего вопросов:", str(len(self.questions))),
            ("⭐ Максимальный балл:", str(self.get_total_points())),
            ("⏱ Время выполнения:", "60 минут")
        ]

        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i] if i < len(info_table.rows) else info_table.add_row()
            label_cell = row.cells[0]
            value_cell = row.cells[1]

            self._apply_cell_formatting(label_cell, label, bold=True, bg_color='F0F0F0')
            self._apply_cell_formatting(value_cell, value)

        if self.description:
            desc_para = self.document.add_paragraph()
            desc_para.add_run(self.description).italic = True

        self.document.add_paragraph("_" * 50)
        self.document.add_paragraph()

    def _apply_cell_formatting(self, cell, text: str, font_size: int = 11,
                               bold: bool = False, italic: bool = False,
                               alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
                               bg_color: str = None):
        """Применяет форматирование к ячейке"""
        cell.text = text
        paragraph = cell.paragraphs[0]
        paragraph.alignment = alignment

        if paragraph.runs:
            run = paragraph.runs[0]
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.italic = italic

        if bg_color:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), bg_color)
            cell._tc.get_or_add_tcPr().append(shading_elm)

    def get_total_points(self) -> int:
        """Возвращает общее количество баллов"""
        return sum(q.points for q in self.questions)

    def generate(self, output_filename: str = "test_questions.docx"):
        """Генерирует документ со всеми вопросами"""
        # Добавляем заголовок
        self.add_header()

        # Добавляем все вопросы
        for i, question in enumerate(self.questions, start=1):
            question.render_to_document(self.document, i)
            # self.document.add_paragraph()
            # Добавляем разделитель между вопросами
            # if i < len(self.questions):
            #     separator = self.document.add_paragraph("•" * 70)
            #     # separator.add_run()
            #     separator.runs[0].font.size = Pt(8)
            #     separator.runs[0].font.color.rgb = RGBColor(192, 192, 192)

        # Добавляем страницу для ответов
        # self.document.add_page_break()
        # self._add_answer_sheet()

        # Сохраняем документ
        self.document.save(output_filename)
        print(f"✅ Тест успешно создан: {output_filename}")
        print(f"📊 Всего вопросов: {len(self.questions)}")
        print(f"⭐ Максимальный балл: {self.get_total_points()}")

    def _add_answer_sheet(self):
        """Добавляет лист ответов"""
        title = self.document.add_heading("📝 ЛИСТ ОТВЕТОВ", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        answer_table = self.document.add_table(rows=len(self.questions) + 1, cols=4)
        answer_table.style = 'Table Grid'
        answer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Заголовки
        headers = ['№ вопроса', 'Тип вопроса', 'Ответ', 'Балл']
        for i, header in enumerate(headers):
            cell = answer_table.rows[0].cells[i]
            self._apply_cell_formatting(cell, header, bold=True,
                                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                       bg_color='D9E1F2')

        # Заполняем строки для ответов
        for i, question in enumerate(self.questions, start=1):
            row = answer_table.rows[i]

            # Номер вопроса
            self._apply_cell_formatting(row.cells[0], str(i),
                                       alignment=WD_ALIGN_PARAGRAPH.CENTER)

            # Тип вопроса
            type_names = {
                QuestionType.SINGLE_CHOICE: "Один ответ",
                QuestionType.MULTIPLE_CHOICE: "Несколько ответов",
                QuestionType.MATCHING: "Сопоставление",
                QuestionType.SORTING: "Сортировка"
            }
            self._apply_cell_formatting(row.cells[1], type_names[question.question_type],
                                       alignment=WD_ALIGN_PARAGRAPH.CENTER)

            # Поле для ответа
            self._apply_cell_formatting(row.cells[2], "_______________",
                                       alignment=WD_ALIGN_PARAGRAPH.CENTER)

            # Поле для балла
            self._apply_cell_formatting(row.cells[3], f"/{question.points}",
                                       alignment=WD_ALIGN_PARAGRAPH.CENTER)

class LabGenerator:
    def __init__(self, number: int, title: str, theory: List[str],
                 template: str = "Расчётно-графическая задача №{}",
                 temp_complete: str = "{}\nрасчётно-графической задачи", doc: Optional[Docx] = None):
        self.document = doc
        if not self.document:
            self.document = Document()
        self.questions = []
        self.name = template.format(number)
        self.title = title
        self.theory = theory
        self.font_name="Times New Roman"
        self.complete = temp_complete

    def _apply_cell_formatting(self, cell, text: str, font_size: int = 11, font_name: str = "Arial",
                               bold: bool = False, italic: bool = False, sub_text: str | None = None,
                               alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
                               bg_color: str = None):
        """Применяет форматирование к ячейке"""
        cell.text = text
        paragraph = cell.paragraphs[0]
        paragraph.alignment = alignment

        if paragraph.runs:
            run = paragraph.runs[0]
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.italic = italic
            run.font.name = font_name

        if bg_color:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), bg_color)
            cell._tc.get_or_add_tcPr().append(shading_elm)

        if sub_text:
            sub_text_paragraph = cell.add_paragraph(sub_text)
            sub_text_paragraph.alignment = alignment
            for run in sub_text_paragraph.runs:
                run.font.size = Pt(font_size - 3)
                run.font.name = self.font_name


    def add_question(self, question: TestQuestion):
        """Добавляет вопрос в тест"""
        self.questions.append(question)

    def add_header(self):
        """Добавляет заголовок теста"""
        # Основной заголовок
        number = self.document.add_heading(self.name, level=1)
        number.paragraph_format.line_spacing = 1
        number.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in number.runs:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = self.font_name
        title = self.document.add_heading(f"«{self.title}»", level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = self.font_name
        par = self.document.add_paragraph()
        par.paragraph_format.line_spacing = 0

    def add_theory(self, name: str = "Теоретическое введение"):
        theory_head = self.document.add_paragraph(name)
        theory_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in theory_head.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = self.font_name
        task = ThermsTask(self.document)
        task.create_tables(self.theory)
        self.document.add_page_break()

    def add_questions_to_lab(self, name: str = "Контрольные вопросы"):
        theory_head = self.document.add_paragraph(name)
        theory_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in theory_head.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = self.font_name
        # Добавляем все вопросы
        for i, question in enumerate(self.questions, start=1):
            question.render_to_document(self.document, i)
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.line_spacing = 0

    def add_control(self):
        self.document.add_paragraph()
        table = self.document.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        for row, text in zip(table.rows, ["Дата выполнения", "Дата защиты"]):
            self._apply_cell_formatting(row.cells[0], self.complete.format(text),
                                        12, self.font_name, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                        sub_text="(заполняется преподавателем)")
            # row.height = Pt(48)

    def generate(self, output_filename: Optional[str] = None):
        """Генерирует документ со всеми вопросами"""
        # Добавляем заголовок
        print(f"📊 Генерируем {self.name}: {self.title}")
        self.add_header()
        print("✅ Добавлен заголовок")
        # Add theory
        self.add_theory()
        print("✅ Добавлена теория")
        # Add questions
        self.add_questions_to_lab(name="Тестовые задания и задачи")
        print("✅ Добавлены тестовые задания")
        # Add control
        self.add_control()

        if output_filename:
            # Сохраняем документ
            self.document.save(output_filename)
            print(f"✅ Лабораторная успешно создана: {output_filename}")
            # print(f"📊 Всего вопросов: {len(self.questions)}")
            # print(f"⭐ Максимальный балл: {self.get_total_points()}")

# ============= ПРИМЕР ИСПОЛЬЗОВАНИЯ =============
def build_labs(project: Project, path_to_doc: str, directory = "images_task"):
    def data_shuffle(task):
        lst = task.answer.correct + task.answer.wrong
        shuffle(lst)
        return lst

    def elem_shuffle(elems):
        shuffle(elems)
        return elems
    doc = Document()
    for i, elem in enumerate(project.elems, 1):
        lab = LabGenerator(i, elem.name, elem.theory, doc=doc)
        for task in elem.tasks:
            if task.question.types == QuestionsType.CHOICE.value and len(task.answer.correct) > 1:
                lab.add_question(
                    MultipleChoiceQuestion(
                    question_text=task.question.text_data,
                    options=data_shuffle(task),
                    correct_answers=[0, 2],  # Django, Flask, FastAPI
                    points=2
                ))
                continue
            match task.question.types:
                case QuestionsType.CHOICE.value:
                    lab.add_question(
                        SingleChoiceQuestion(
                            question_text=task.question.text_data,
                            options=data_shuffle(task),
                            correct_answer=2,  # Python (индексация с 0)
                            points=1
                        )
                    )
                case QuestionsType.MATCHING.value:
                    lab.add_question(
                        MatchingQuestion(
                            question_text=task.question.text_data,
                            left_items=elem_shuffle(task.answer.first),
                            right_items=elem_shuffle(task.answer.second),
                            matches=[(0, 1)],  # Python → Data Science, JavaScript → Web, etc.
                            points=3
                        )
                    )
                case QuestionsType.SORTING.value:
                    lab.add_question(
                        SortingQuestion(
                            question_text=task.question.text_data,
                            items=elem_shuffle(task.answer.steps),
                            correct_order=[1, 2, 3, 0],  # Quick → Merge → Selection → Bubble
                            points=2
                        )
                    )
                case QuestionsType.NUMBER.value:
                    path = None
                    if task.question.image:
                        response = requests.get(task.question.image)
                        if response.status_code == 200:
                            if not os.path.exists(directory):
                                os.makedirs(directory)
                                print("Директория создана")

                            path = os.path.join(directory, task.question.image.split('/')[-1])
                            with open(path, 'wb') as file:
                                file.write(response.content)
                            print('Изображение сохранено', path)
                        else:
                            print('Ошибка загрузки изображения')
                    lab.add_question(
                        NumberTask(
                            question_text=task.question.text_data,
                            image_path=path
                        )
                    )
        lab.generate()
        doc.add_page_break()

    doc.save(path_to_doc)
    print("\n✨ Проект успешно создан! Откройте файл:", path_to_doc)


def create_labs():
    lab = LabGenerator(1, 'Physical tehemes', ['Theormex', 'Phys', 'Datalens'])
    questions = [
        SingleChoiceQuestion(
            question_text="Какой язык программирования наиболее популярен для Data Science?",
            options=["Java", "C++", "Python", "JavaScript"],
            correct_answer=2,  # Python (индексация с 0)
            points=1
        ),
        MultipleChoiceQuestion(
            question_text="Какие из перечисленных являются фреймворками Python?",
            options=["Django", "React", "Flask", "Spring", "FastAPI"],
            correct_answers=[0, 2, 4],  # Django, Flask, FastAPI
            points=2
        ),
        MatchingQuestion(
            question_text="Сопоставьте язык программирования с его основным применением:",
            left_items=["Python", "JavaScript", "SQL", "HTML/CSS"],
            right_items=["Веб-разработка", "Анализ данных и AI", "Работа с базами данных", "Структура веб-страниц"],
            matches=[(0, 1), (1, 0), (2, 2), (3, 3)],  # Python → Data Science, JavaScript → Web, etc.
            points=3
        ),
        SortingQuestion(
            question_text="Расположите алгоритмы сортировки от наилучшей к наихудшей временной сложности (в среднем):",
            items=["Сортировка пузырьком", "Быстрая сортировка", "Сортировка слиянием", "Сортировка выбором"],
            correct_order=[1, 2, 3, 0],  # Quick → Merge → Selection → Bubble
            points=2
        ),
        NumberTask(
            question_text="Определите ускорение бруска массой 2 кг, если на него действует сила 10 Н. Трением пренебречь.",
            image_path="C:\\Users\\Юрий Солдатов\\Pictures\\photo_2025-03-08_12-21-06.jpg"
        )
    ]
    for elem in questions:
        lab.add_question(elem)

    lab.generate("комплексный_тест.docx")

def create_sample_test():
    """Создает пример теста с разными типами вопросов"""

    # Создаем генератор тестов
    test_generator = TestGenerator(
        title="📚 ИТОГОВОЕ ТЕСТИРОВАНИЕ",
        description="Пожалуйста, внимательно прочитайте каждый вопрос и выберите правильный ответ. "
                    "Время выполнения: 60 минут. Удачи!"
    )

    # 1. Вопрос с одним правильным ответом
    question1 = SingleChoiceQuestion(
        question_text="Какой язык программирования наиболее популярен для Data Science?",
        options=["Java", "C++", "Python", "JavaScript"],
        correct_answer=2,  # Python (индексация с 0)
        points=1
    )
    test_generator.add_question(question1)

    # 2. Вопрос с несколькими правильными ответами
    question2 = MultipleChoiceQuestion(
        question_text="Какие из перечисленных являются фреймворками Python?",
        options=["Django", "React", "Flask", "Spring", "FastAPI"],
        correct_answers=[0, 2, 4],  # Django, Flask, FastAPI
        points=2
    )
    test_generator.add_question(question2)

    # 3. Вопрос на сопоставление
    question3 = MatchingQuestion(
        question_text="Сопоставьте язык программирования с его основным применением:",
        left_items=["Python", "JavaScript", "SQL", "HTML/CSS"],
        right_items=["Веб-разработка", "Анализ данных и AI", "Работа с базами данных", "Структура веб-страниц"],
        matches=[(0, 1), (1, 0), (2, 2), (3, 3)],  # Python → Data Science, JavaScript → Web, etc.
        points=3
    )
    test_generator.add_question(question3)

    # 4. Вопрос на сортировку
    question4 = SortingQuestion(
        question_text="Расположите этапы разработки ПО в правильном порядке:",
        items=["Тестирование", "Анализ требований", "Разработка", "Проектирование", "Внедрение"],
        correct_order=[1, 3, 2, 0, 4],  # Анализ → Проектирование → Разработка → Тестирование → Внедрение
        points=2
    )
    test_generator.add_question(question4)

    # 5. Еще один вопрос с одним ответом
    question5 = SingleChoiceQuestion(
        question_text="Что означает аббревиатура SQL?",
        options=["Structured Query Language", "Simple Query Language", "Standard Question Language", "System Query Logic"],
        correct_answer=0,
        points=1
    )
    test_generator.add_question(question5)

    # 6. Вопрос с несколькими ответами о базах данных
    question6 = MultipleChoiceQuestion(
        question_text="Какие из следующих СУБД являются реляционными?",
        options=["MySQL", "MongoDB", "PostgreSQL", "Redis", "Oracle Database"],
        correct_answers=[0, 2, 4],  # MySQL, PostgreSQL, Oracle
        points=2
    )
    test_generator.add_question(question6)

    # 7. Сопоставление терминов
    question7 = MatchingQuestion(
        question_text="Сопоставьте понятия с их определениями:",
        left_items=["API", "IDE", "ORM", "REST"],
        right_items=["Среда разработки", "Объектно-реляционное отображение", "Архитектурный стиль API", "Интерфейс программирования"],
        matches=[(0, 3), (1, 0), (2, 1), (3, 2)],
        points=3
    )
    test_generator.add_question(question7)

    # 8. Сортировка алгоритмов
    question8 = SortingQuestion(
        question_text="Расположите алгоритмы сортировки от наилучшей к наихудшей временной сложности (в среднем):",
        items=["Сортировка пузырьком", "Быстрая сортировка", "Сортировка слиянием", "Сортировка выбором"],
        correct_order=[1, 2, 3, 0],  # Quick → Merge → Selection → Bubble
        points=2
    )
    test_generator.add_question(question8)

    # Генерируем документ
    test_generator.generate("комплексный_тест.docx")

    return test_generator


def create_custom_test(questions_data: List[dict], title: str, description: str = ""):
    """
    Создает тест на основе переданных данных

    questions_data: список словарей с данными вопросов
    Каждый словарь должен содержать:
        - type: 'single', 'multiple', 'matching', 'sorting'
        - text: текст вопроса
        - points: баллы (опционально)
        - и специфичные для типа поля
    """
    test_generator = TestGenerator(title, description)

    for q_data in questions_data:
        q_type = q_data.get('type')
        points = q_data.get('points', 1)

        if q_type == 'single':
            question = SingleChoiceQuestion(
                question_text=q_data['text'],
                options=q_data['options'],
                correct_answer=q_data['correct_answer'],
                points=points
            )
        elif q_type == 'multiple':
            question = MultipleChoiceQuestion(
                question_text=q_data['text'],
                options=q_data['options'],
                correct_answers=q_data['correct_answers'],
                points=points
            )
        elif q_type == 'matching':
            question = MatchingQuestion(
                question_text=q_data['text'],
                left_items=q_data['left_items'],
                right_items=q_data['right_items'],
                matches=q_data['matches'],
                points=points
            )
        elif q_type == 'sorting':
            question = SortingQuestion(
                question_text=q_data['text'],
                items=q_data['items'],
                correct_order=q_data['correct_order'],
                points=points
            )
        else:
            raise ValueError(f"Неизвестный тип вопроса: {q_type}")

        test_generator.add_question(question)

    test_generator.generate(f"{title.lower().replace(' ', '_')}.docx")
    return test_generator


if __name__ == "__main__":
    # Создаем пример теста
    print("🔨 Создание тестового задания...")
    # test = create_sample_test()

    # Пример создания кастомного теста
    # custom_questions = [
    #     {
    #         'type': 'single',
    #         'text': 'Сколько бит в одном байте?',
    #         'options': ['4', '8', '16', '32'],
    #         'correct_answer': 1,
    #         'points': 1
    #     },
    #     {
    #         'type': 'multiple',
    #         'text': 'Какие операционные системы являются свободными (open source)?',
    #         'options': ['Windows', 'Linux', 'macOS', 'FreeBSD', 'Android'],
    #         'correct_answers': [1, 3, 4],
    #         'points': 2
    #     },
    #     {
    #         'type': 'sorting',
    #         'text': 'Расположите числа в порядке возрастания:',
    #         'items': ['42', '7', '15', '3', '99'],
    #         'correct_order': [3, 1, 2, 0, 4],  # 3, 7, 15, 42, 99
    #         'points': 1
    #     }
    # ]

    # Раскомментируйте для создания кастомного теста
    # custom_test = create_custom_test(custom_questions, "Кастомный тест", "Тест по информатике")
    create_labs()
    print("\n✨ Лабораторная успешно создана! Откройте файл 'комплексный_тест.docx'")